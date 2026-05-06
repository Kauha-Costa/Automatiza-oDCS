import re
import os
import copy
from docx import Document
from docx.oxml.ns import qn


PADRAO_DATA = re.compile(r"\b\d{2}/\d{2}/\d{4}\b")


def _substituir_em_run(run, nova_data: str) -> bool:
    """Substitui datas em um run preservando toda a formatação original."""
    if PADRAO_DATA.search(run.text):
        run.text = PADRAO_DATA.sub(nova_data, run.text)
        return True
    return False


def _runs_do_paragrafo_tem_data(paragrafo) -> bool:
    return bool(PADRAO_DATA.search(paragrafo.text))


def atualizar_data_certificado(caminho_arquivo: str, nova_data: str) -> bool:
    """
    Atualiza a data em um arquivo .docx preservando formatação por run.
    Retorna True se alguma data foi encontrada e substituída.
    """
    doc = Document(caminho_arquivo)
    atualizado = False

    # Percorre parágrafos normais
    for paragrafo in doc.paragraphs:
        if not _runs_do_paragrafo_tem_data(paragrafo):
            continue

        texto_original = paragrafo.text

        # Caso simples: parágrafo tem apenas 1 run ou runs sem quebra de data
        substituiu_por_run = False
        for run in paragrafo.runs:
            if _substituir_em_run(run, nova_data):
                substituiu_por_run = True
                atualizado = True

        # Fallback: a data pode estar fragmentada entre runs (ex: "01" + "/" + "05" + "/" + "2025")
        # Reconstrói o parágrafo se a data ainda aparecer no texto original mas não foi substituída
        if not substituiu_por_run and PADRAO_DATA.search(texto_original):
            # Concatena o texto de todos os runs, substitui, redistribui
            texto_novo = PADRAO_DATA.sub(nova_data, texto_original)
            if texto_novo != texto_original:
                # Coloca tudo no primeiro run e limpa os demais
                if paragrafo.runs:
                    paragrafo.runs[0].text = texto_novo
                    for run in paragrafo.runs[1:]:
                        run.text = ""
                atualizado = True

    # Percorre tabelas também (certificados podem ter datas em tabelas)
    for tabela in doc.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                for paragrafo in celula.paragraphs:
                    if not _runs_do_paragrafo_tem_data(paragrafo):
                        continue
                    texto_original = paragrafo.text
                    substituiu_por_run = False
                    for run in paragrafo.runs:
                        if _substituir_em_run(run, nova_data):
                            substituiu_por_run = True
                            atualizado = True
                    if not substituiu_por_run and PADRAO_DATA.search(texto_original):
                        texto_novo = PADRAO_DATA.sub(nova_data, texto_original)
                        if texto_novo != texto_original:
                            if paragrafo.runs:
                                paragrafo.runs[0].text = texto_novo
                                for run in paragrafo.runs[1:]:
                                    run.text = ""
                            atualizado = True

    try:
        doc.save(caminho_arquivo)
    except PermissionError:
        raise PermissionError(
            f"Não foi possível salvar '{os.path.basename(caminho_arquivo)}'. "
            f"Feche o Word antes de executar."
        )

    return atualizado


def atualizar_certificados_em_pasta(pasta: str, nova_data: str, log_callback=print, cancelar_flag=None):
    """
    Atualiza todos os .docx na pasta fornecida.
    cancelar_flag: lista com um bool [False]; se virar [True], interrompe.
    """
    if not os.path.exists(pasta):
        log_callback(f"❌ Pasta não encontrada: {pasta}")
        return 0, 0

    arquivos = [arq for arq in os.listdir(pasta) if arq.lower().endswith(".docx")]

    if not arquivos:
        log_callback("⚠️  Nenhum arquivo .docx encontrado na pasta.")
        return 0, 0

    total = len(arquivos)
    atualizados = 0
    sem_data = 0

    for i, arquivo in enumerate(arquivos):
        if cancelar_flag and cancelar_flag[0]:
            log_callback("⛔ Operação cancelada pelo usuário.")
            break

        caminho = os.path.join(pasta, arquivo)
        log_callback(f"⏳ [{i+1}/{total}] Atualizando: {arquivo}")

        try:
            foi_atualizado = atualizar_data_certificado(caminho, nova_data)
            if foi_atualizado:
                log_callback(f"✅ Atualizado: {arquivo}")
                atualizados += 1
            else:
                log_callback(f"ℹ️  Sem data encontrada: {arquivo}")
                sem_data += 1
        except PermissionError as e:
            log_callback(f"❌ {e}")
        except Exception as e:
            log_callback(f"❌ Erro em '{arquivo}': {e}")

    return atualizados, sem_data


def listar_certificados(pasta: str) -> list[str]:
    """Retorna lista de arquivos .docx na pasta."""
    if not os.path.exists(pasta):
        return []
    return sorted([arq for arq in os.listdir(pasta) if arq.lower().endswith(".docx")])
